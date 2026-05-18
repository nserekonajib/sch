# competenceReportCard.py - Complete Competence-Based Report Card System
from flask import Blueprint, render_template, request, jsonify, session, send_file
from supabase import create_client, Client
import os
import uuid
from datetime import datetime
import json
import re
from functools import wraps
from dotenv import load_dotenv
import io
import requests
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Inches
import tempfile
import subprocess
import sys
import zipfile
from routes.auth.auth import role_required
from routes.accounts.accounts import get_institute_id as get_institute_id_func

load_dotenv()

# Initialize Supabase client
SUPABASE_URL = os.getenv('SUPABASE_URL')
SUPABASE_KEY = os.getenv('SUPABASE_KEY')
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

competence_bp = Blueprint('competence', __name__, url_prefix='/competence-report')

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return jsonify({'success': False, 'message': 'Please login'}), 401
        return f(*args, **kwargs)
    return decorated_function

def parse_exam_name(exam_name):
    patterns = {
        'A': re.compile(r'^A\d+$', re.IGNORECASE),
        'BOT': re.compile(r'^BOT$', re.IGNORECASE),
        'EOT': re.compile(r'^EOT$', re.IGNORECASE),
        'MT': re.compile(r'^MT$', re.IGNORECASE),
    }
    for pattern_type, pattern in patterns.items():
        if pattern.match(exam_name.strip().upper()):
            return pattern_type
    return 'OTHER'

def convert_docx_to_pdf_with_fallback(input_path, output_dir, output_filename=None):
    """
    Convert DOCX to PDF with two methods:
    1. Try docx2pdf first (Windows with Microsoft Word)
    2. Fallback to LibreOffice (cross-platform)
    """
    temp_files = []
    
    # Ensure output_dir exists
    os.makedirs(output_dir, exist_ok=True)
    
    # Determine output PDF path
    if output_filename:
        pdf_path = os.path.join(output_dir, output_filename)
    else:
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
    
    # Method 1: Try docx2pdf (Windows with MS Word)
    try:
        print(f"Attempt 1: Converting {input_path} to PDF using docx2pdf...")
        from docx2pdf import convert
        convert(input_path, pdf_path)
        
        # Check if conversion was successful
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            print(f"✓ PDF conversion successful with docx2pdf! Size: {os.path.getsize(pdf_path)} bytes")
            return pdf_path
        else:
            raise Exception("docx2pdf produced empty or missing file")
            
    except ImportError:
        print("docx2pdf not installed, trying LibreOffice...")
    except Exception as e:
        print(f"docx2pdf conversion failed: {e}, trying LibreOffice...")
    
    # Method 2: Try LibreOffice (cross-platform)
    try:
        print(f"Attempt 2: Converting {input_path} to PDF using LibreOffice...")
        
        # Check if soffice is available
        soffice_paths = [
            "soffice",  # Default in PATH
            "/usr/bin/soffice",  # Linux
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",  # Windows
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",  # Windows 32-bit
        ]
        
        soffice_cmd = None
        for path in soffice_paths:
            if os.path.exists(path) or (path == "soffice" and subprocess.run(["which", "soffice"], capture_output=True).returncode == 0):
                soffice_cmd = path
                break
        
        if not soffice_cmd:
            raise Exception("LibreOffice (soffice) not found on system")
        
        # Run LibreOffice conversion
        result = subprocess.run([
            soffice_cmd,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            input_path
        ], capture_output=True, text=True, timeout=60)
        
        # Check if PDF was created (LibreOffice saves with original name)
        expected_pdf = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(input_path))[0]}.pdf")
        
        if os.path.exists(expected_pdf) and os.path.getsize(expected_pdf) > 0:
            # Rename to desired output path if different
            if expected_pdf != pdf_path:
                os.rename(expected_pdf, pdf_path)
            print(f"✓ PDF conversion successful with LibreOffice! Size: {os.path.getsize(pdf_path)} bytes")
            return pdf_path
        else:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")
            
    except subprocess.TimeoutExpired:
        print("LibreOffice conversion timed out after 60 seconds")
        raise Exception("PDF conversion timed out")
    except Exception as e:
        print(f"LibreOffice conversion failed: {e}")
        raise Exception(f"All PDF conversion methods failed: {str(e)}")

def merge_pdfs(pdf_paths, output_path):
    """
    Merge multiple PDF files into a single PDF
    """
    try:
        from PyPDF2 import PdfMerger
        merger = PdfMerger()
        
        for pdf_path in pdf_paths:
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                merger.append(pdf_path)
            else:
                print(f"Warning: PDF {pdf_path} is missing or empty")
        
        merger.write(output_path)
        merger.close()
        print(f"✓ Merged {len(pdf_paths)} PDFs into {output_path}")
        return True
        
    except ImportError:
        print("PyPDF2 not installed, trying pdfkit alternative...")
        # Alternative: Use pdfunite (Linux command line)
        try:
            cmd = ["pdfunite"] + pdf_paths + [output_path]
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0:
                print(f"✓ Merged PDFs using pdfunite")
                return True
            else:
                raise Exception("pdfunite failed")
        except Exception as e:
            print(f"PDF merge failed: {e}")
            return False

@competence_bp.route('/')
@role_required(['owner', 'teacher', 'accountant'])
def index():
    user = session.get('user')
    institute_id = get_institute_id_func(user['id'])
    
    if not institute_id:
        return render_template('competence/index.html', exams=[], classes=[], students=[], institute=None)
    
    try:
        # Get institute details
        institute_response = supabase.table('institutes')\
            .select('*')\
            .eq('id', institute_id)\
            .execute()
        
        institute = institute_response.data[0] if institute_response.data else {}
        
        exams_response = supabase.table('exams')\
            .select('*')\
            .eq('institute_id', institute_id)\
            .order('exam_date', desc=True)\
            .execute()
        
        exams = exams_response.data if exams_response.data else []
        
        classes_response = supabase.table('classes')\
            .select('*')\
            .eq('institute_id', institute_id)\
            .order('name')\
            .execute()
        
        classes = classes_response.data if classes_response.data else []
        
        students_response = supabase.table('students')\
            .select('id, name, student_id, class_id, classes(name), photo_url, gender, category')\
            .eq('institute_id', institute_id)\
            .eq('status', 'active')\
            .order('name')\
            .execute()
        
        students = students_response.data if students_response.data else []
        
        return render_template('competence/index.html', exams=exams, classes=classes, students=students, institute=institute)
        
    except Exception as e:
        print(f"Error loading competence page: {e}")
        return render_template('competence/index.html', exams=[], classes=[], students=[], institute=None)

@competence_bp.route('/api/exams/classify', methods=['POST'])
@role_required(['owner', 'teacher', 'accountant'])
def classify_exams():
    user = session.get('user')
    institute_id = get_institute_id_func(user['id'])
   
    if not institute_id:
        return jsonify({'success': False, 'message': 'Institute not found'}), 400
    
    try:
        data = request.get_json()
        exam_ids = data.get('exam_ids', [])
        
        if not exam_ids:
            return jsonify({'success': False, 'message': 'No exams selected'}), 400
        
        exams_response = supabase.table('exams')\
            .select('*')\
            .in_('id', exam_ids)\
            .execute()
        
        exams = exams_response.data if exams_response.data else []
        
        classified = {'a_series': [], 'bot': None, 'eot': None, 'mt': None, 'other': []}
        
        for exam in exams:
            exam_type = parse_exam_name(exam['exam_name'])
            if exam_type == 'A':
                match = re.search(r'(\d+)', exam['exam_name'])
                number = int(match.group(1)) if match else 0
                classified['a_series'].append({
                    'id': exam['id'], 'name': exam['exam_name'], 
                    'total_marks': exam['total_marks'], 'number': number
                })
            elif exam_type == 'BOT':
                classified['bot'] = {'id': exam['id'], 'name': exam['exam_name'], 'total_marks': exam['total_marks']}
            elif exam_type == 'EOT':
                classified['eot'] = {'id': exam['id'], 'name': exam['exam_name'], 'total_marks': exam['total_marks']}
            elif exam_type == 'MT':
                classified['mt'] = {'id': exam['id'], 'name': exam['exam_name'], 'total_marks': exam['total_marks']}
            else:
                classified['other'].append({'id': exam['id'], 'name': exam['exam_name'], 'total_marks': exam['total_marks']})
        
        classified['a_series'].sort(key=lambda x: x['number'])
        
        return jsonify({'success': True, 'classified': classified})
        
    except Exception as e:
        print(f"Error classifying exams: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

@competence_bp.route('/generate', methods=['POST'])
@role_required(['owner', 'teacher', 'accountant'])
def generate_report():
    user = session.get('user')
    institute_id = get_institute_id_func(user['id'])
    
    if not institute_id:
        return jsonify({'success': False, 'message': 'Institute not found'}), 400
    
    # Store temp files for cleanup
    temp_files = []
    
    try:
        data = request.get_json()
        student_id = data.get('student_id')
        exam_ids = data.get('exam_ids', [])
        term = data.get('term', '')
        year = data.get('year', datetime.now().year)
        format_type = data.get('format', 'docx')
        
        if not student_id or not exam_ids or not term:
            return jsonify({'success': False, 'message': 'Missing required fields'}), 400
        
        # Get institute details
        institute_response = supabase.table('institutes')\
            .select('*')\
            .eq('id', institute_id)\
            .execute()
        
        institute = institute_response.data[0] if institute_response.data else {}
        
        # Get student details
        student_response = supabase.table('students')\
            .select('*, classes(name)')\
            .eq('id', student_id)\
            .eq('institute_id', institute_id)\
            .execute()
        
        if not student_response.data:
            return jsonify({'success': False, 'message': 'Student not found'}), 404
        
        student = student_response.data[0]
        
        # Get exams
        exams_response = supabase.table('exams')\
            .select('*')\
            .in_('id', exam_ids)\
            .execute()
        
        exams = exams_response.data if exams_response.data else []
        
        # Classify exams
        a_series = []
        eot_exam = None
        for exam in exams:
            exam_type = parse_exam_name(exam['exam_name'])
            if exam_type == 'A':
                match = re.search(r'(\d+)', exam['exam_name'])
                number = int(match.group(1)) if match else 0
                a_series.append({
                    'id': exam['id'], 
                    'name': exam['exam_name'], 
                    'total_marks': exam['total_marks'], 
                    'number': number
                })
            elif exam_type == 'EOT':
                eot_exam = exam
        
        a_series.sort(key=lambda x: x['number'])
        
        # Get subjects
        class_id = student.get('class_id')
        subjects_response = supabase.table('class_subjects')\
            .select('*, subjects(name)')\
            .eq('class_id', class_id)\
            .eq('institute_id', institute_id)\
            .execute()
        
        subjects = subjects_response.data if subjects_response.data else []
        
        # Get marks for each exam
        all_marks = {}
        for exam in exams:
            marks_response = supabase.table('exam_marks')\
                .select('*')\
                .eq('exam_id', exam['id'])\
                .eq('student_id', student_id)\
                .eq('institute_id', institute_id)\
                .execute()
            
            marks_dict = {}
            for mark in marks_response.data if marks_response.data else []:
                marks_dict[mark['subject_id']] = float(mark['obtained_marks'])
            all_marks[exam['id']] = marks_dict
        
        # Prepare subject data
        subject_results = []
        for idx, subject in enumerate(subjects):
            subject_name = subject['subjects']['name'] if subject.get('subjects') else 'N/A'
            
            a_series_marks = []
            for a_exam in a_series:
                obtained = all_marks.get(a_exam['id'], {}).get(subject['subject_id'])
                a_series_marks.append(obtained if obtained is not None else 0)
            
            while len(a_series_marks) < 3:
                a_series_marks.append(0)
            
            a_series_avg = sum(a_series_marks) / len(a_series_marks) if a_series_marks else 0
            
            if a_series and a_series[0]['total_marks'] > 0:
                twenty_percent = (a_series_avg / a_series[0]['total_marks']) * 20
            else:
                twenty_percent = 0
            
            if eot_exam and eot_exam.get('total_marks', 0) > 0:
                eot_mark = all_marks.get(eot_exam['id'], {}).get(subject['subject_id'], 0)
                eighty_percent = (eot_mark / eot_exam['total_marks']) * 80
            else:
                eot_mark = 0
                eighty_percent = 0
            
            total_percentage = twenty_percent + eighty_percent
            
            identifier = '1' if total_percentage < 40 else ('2' if total_percentage < 70 else '3')
            
            if total_percentage >= 80:
                grade = 'A'
                remark = 'Exceptional'
            elif total_percentage >= 70:
                grade = 'B'
                remark = 'Outstanding'
            elif total_percentage >= 60:
                grade = 'C'
                remark = 'Satisfactory'
            elif total_percentage >= 40:
                grade = 'D'
                remark = 'Basic'
            else:
                grade = 'E'
                remark = 'Insufficient'
            
            subject_results.append({
                'name': subject_name,
                'a1': f"{a_series_marks[0]:.1f}",
                'a2': f"{a_series_marks[1]:.1f}",
                'a3': f"{a_series_marks[2]:.1f}",
                'avg': f"{a_series_avg:.1f}",
                'twenty': f"{twenty_percent:.1f}",
                'eighty': f"{eighty_percent:.1f}",
                'total': f"{total_percentage:.1f}",
                'identifier': identifier,
                'grade': grade,
                'remark': remark
            })
        
        if subject_results:
            overall_percentage = sum([float(s['total']) for s in subject_results]) / len(subject_results)
        else:
            overall_percentage = 0
        
        if overall_percentage >= 80:
            overall_grade = 'A'
            overall_achievement = 'EXCELLENT'
        elif overall_percentage >= 70:
            overall_grade = 'B'
            overall_achievement = 'VERY GOOD'
        elif overall_percentage >= 60:
            overall_grade = 'C'
            overall_achievement = 'GOOD'
        elif overall_percentage >= 40:
            overall_grade = 'D'
            overall_achievement = 'AVERAGE'
        else:
            overall_grade = 'E'
            overall_achievement = 'NEEDS IMPROVEMENT'
        
        overall_ident = '3' if overall_percentage >= 70 else ('2' if overall_percentage >= 40 else '1')
        
        def safe_upper(value, default='N/A'):
            if value is None:
                return default
            return str(value).upper()
        
        def safe_get(dict_obj, key, default=''):
            value = dict_obj.get(key, default)
            if value is None:
                return default
            return value
        
        # Load the Word template
        template_dir = os.path.join(os.path.dirname(__file__), 'templates')
        template_path = os.path.join(template_dir, 'competence_report_template.docx')
        
        if not os.path.exists(template_path):
            os.makedirs(template_dir, exist_ok=True)
            create_default_template(template_path)
        
        doc = DocxTemplate(template_path)
        
        # Handle images (same as before)
        logo_url = institute.get('logo_url', '')
        context_logo = ""
        if logo_url and logo_url.startswith(('http://', 'https://')):
            try:
                response = requests.get(logo_url, timeout=10)
                if response.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_logo:
                        tmp_logo.write(response.content)
                        tmp_logo_path = tmp_logo.name
                        temp_files.append(tmp_logo_path)
                        context_logo = InlineImage(doc, tmp_logo_path, width=Mm(25))
            except Exception as e:
                print(f"Error loading logo: {e}")
        
        photo_url = student.get('photo_url', '')
        context_photo = ""
        if photo_url and photo_url.startswith(('http://', 'https://')):
            try:
                response = requests.get(photo_url, timeout=10)
                if response.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_photo:
                        tmp_photo.write(response.content)
                        tmp_photo_path = tmp_photo.name
                        temp_files.append(tmp_photo_path)
                        context_photo = InlineImage(doc, tmp_photo_path, width=Mm(25))
            except Exception as e:
                print(f"Error loading photo: {e}")
        
        # Prepare context
        context = {
            'logo': context_logo,
            'student_photo': context_photo,
            'SCHOOL_NAME': safe_upper(institute.get('institute_name'), 'ST CHARLES LWANGA SS-AKASHANDA'),
            'ADDRESS': safe_get(institute, 'address', 'Kampala, Uganda'),
            'PHONE_NUMBER': safe_get(institute, 'phone_number', '+256757632756'),
            'EMAIL': safe_get(institute, 'email', 'info@school.ac.ug'),
            'TERM': safe_upper(term, term),
            'YEAR': str(year),
            'STUDENT_NAME': safe_upper(student.get('name'), 'STUDENT'),
            'GENDER': safe_upper(student.get('gender'), 'NOT SPECIFIED'),
            'CLASS_NAME': safe_upper(student['classes']['name']) if student.get('classes') else 'N/A',
            'SECTION': safe_upper(student.get('category'), 'DAY'),
            'subjects': subject_results,
            'overall_avg': f"{overall_percentage:.1f}",
            'overall_ident': overall_ident,
            'overall_achievement': overall_achievement,
            'overall_grade': overall_grade,
            'teacher_comment': f"{safe_get(student, 'name', 'The student')} is making good progress.",
            'next_term_date': f"05/{datetime.now().year + 1 if datetime.now().month > 8 else datetime.now().month}/2026",
            'printed_date': datetime.now().strftime('%d/%m/%Y'),
        }
        
        # Render template
        doc.render(context)
        
        # Save to bytes buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Handle PDF conversion if requested
        if format_type == 'pdf':
            try:
                # Save DOCX to temp file
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                    tmp_docx.write(docx_buffer.getvalue())
                    docx_path = tmp_docx.name
                    temp_files.append(docx_path)
                
                # Convert with fallback
                pdf_path = convert_docx_to_pdf_with_fallback(docx_path, tempfile.gettempdir())
                temp_files.append(pdf_path)
                
                # Read PDF
                with open(pdf_path, 'rb') as f:
                    pdf_buffer = io.BytesIO(f.read())
                
                return send_file(
                    pdf_buffer,
                    as_attachment=True,
                    download_name=f"competence_report_{student['name']}_{term}_{year}.pdf",
                    mimetype='application/pdf'
                )
            except Exception as e:
                print(f"PDF conversion failed: {e}, falling back to DOCX")
                docx_buffer.seek(0)
                return send_file(
                    docx_buffer,
                    as_attachment=True,
                    download_name=f"competence_report_{student['name']}_{term}_{year}.docx",
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
        
        # Return DOCX by default
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=f"competence_report_{student['name']}_{term}_{year}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Error generating competence report: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500
    
    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                print(f"Warning: Could not delete {temp_file}: {e}")

@competence_bp.route('/generate-class', methods=['POST'])
@role_required(['owner', 'teacher', 'accountant'])
def generate_class_reports():
    """
    Generate competence reports for all students in a class
    Returns a single merged PDF containing all student reports
    """
    user = session.get('user')
    institute_id = get_institute_id_func(user['id'])
    
    if not institute_id:
        return jsonify({'success': False, 'message': 'Institute not found'}), 400
    
    temp_files = []
    pdf_files = []
    
    try:
        data = request.get_json()
        class_id = data.get('class_id')
        exam_ids = data.get('exam_ids', [])
        term = data.get('term', '')
        year = data.get('year', datetime.now().year)
        
        if not class_id or not exam_ids or not term:
            return jsonify({'success': False, 'message': 'Missing required fields'}), 400
        
        # Get all active students in the class
        students_response = supabase.table('students')\
            .select('*')\
            .eq('class_id', class_id)\
            .eq('institute_id', institute_id)\
            .eq('status', 'active')\
            .order('name')\
            .execute()
        
        students = students_response.data if students_response.data else []
        
        if not students:
            return jsonify({'success': False, 'message': 'No students found in this class'}), 404
        
        print(f"Generating reports for {len(students)} students...")
        
        # Create a temporary directory for individual reports
        temp_dir = tempfile.mkdtemp()
        temp_files.append(temp_dir)
        
        # Generate report for each student
        for idx, student in enumerate(students):
            print(f"Processing {idx+1}/{len(students)}: {student.get('name')}")
            
            # Generate single student report
            report_result = generate_single_student_report(
                student=student,
                institute_id=institute_id,
                exam_ids=exam_ids,
                term=term,
                year=year,
                temp_dir=temp_dir
            )
            
            if report_result:
                pdf_files.append(report_result)
        
        if not pdf_files:
            return jsonify({'success': False, 'message': 'No reports were generated'}), 500
        
        # Merge all PDFs into a single PDF
        merged_pdf_path = os.path.join(temp_dir, f"class_reports_{class_id}_term{term}_{year}.pdf")
        
        if len(pdf_files) == 1:
            # Only one student, just rename the file
            os.rename(pdf_files[0], merged_pdf_path)
        else:
            # Merge multiple PDFs
            merge_success = merge_pdfs(pdf_files, merged_pdf_path)
            if not merge_success:
                # If merge fails, return as ZIP
                return create_zip_archive(pdf_files, class_id, term, year)
        
        # Send the merged PDF
        if os.path.exists(merged_pdf_path) and os.path.getsize(merged_pdf_path) > 0:
            # Get class name
            class_response = supabase.table('classes')\
                .select('name')\
                .eq('id', class_id)\
                .execute()
            class_name = class_response.data[0]['name'] if class_response.data else f"class_{class_id}"
            
            with open(merged_pdf_path, 'rb') as f:
                pdf_buffer = io.BytesIO(f.read())
            
            return send_file(
                pdf_buffer,
                as_attachment=True,
                download_name=f"competence_reports_{class_name}_term{term}_{year}.pdf",
                mimetype='application/pdf'
            )
        else:
            return jsonify({'success': False, 'message': 'Failed to create merged PDF'}), 500
            
    except Exception as e:
        print(f"Error generating class reports: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500
    
    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    if os.path.isdir(temp_file):
                        import shutil
                        shutil.rmtree(temp_file)
                    else:
                        os.unlink(temp_file)
            except Exception as e:
                print(f"Warning: Could not delete {temp_file}: {e}")

def generate_single_student_report(student, institute_id, exam_ids, term, year, temp_dir):
    """
    Generate a single student report and return the PDF path
    """
    try:
        # Get institute details
        institute_response = supabase.table('institutes')\
            .select('*')\
            .eq('id', institute_id)\
            .execute()
        institute = institute_response.data[0] if institute_response.data else {}
        
        # Get exams
        exams_response = supabase.table('exams')\
            .select('*')\
            .in_('id', exam_ids)\
            .execute()
        exams = exams_response.data if exams_response.data else []
        
        # Classify exams
        a_series = []
        eot_exam = None
        for exam in exams:
            exam_type = parse_exam_name(exam['exam_name'])
            if exam_type == 'A':
                match = re.search(r'(\d+)', exam['exam_name'])
                number = int(match.group(1)) if match else 0
                a_series.append({
                    'id': exam['id'], 
                    'name': exam['exam_name'], 
                    'total_marks': exam['total_marks'], 
                    'number': number
                })
            elif exam_type == 'EOT':
                eot_exam = exam
        a_series.sort(key=lambda x: x['number'])
        
        # Get subjects
        class_id = student.get('class_id')
        subjects_response = supabase.table('class_subjects')\
            .select('*, subjects(name)')\
            .eq('class_id', class_id)\
            .eq('institute_id', institute_id)\
            .execute()
        subjects = subjects_response.data if subjects_response.data else []
        
        # Get marks for each exam
        all_marks = {}
        for exam in exams:
            marks_response = supabase.table('exam_marks')\
                .select('*')\
                .eq('exam_id', exam['id'])\
                .eq('student_id', student['id'])\
                .eq('institute_id', institute_id)\
                .execute()
            
            marks_dict = {}
            for mark in marks_response.data if marks_response.data else []:
                marks_dict[mark['subject_id']] = float(mark['obtained_marks'])
            all_marks[exam['id']] = marks_dict
        
        # Prepare subject data
        subject_results = []
        for subject in subjects:
            subject_name = subject['subjects']['name'] if subject.get('subjects') else 'N/A'
            
            a_series_marks = []
            for a_exam in a_series:
                obtained = all_marks.get(a_exam['id'], {}).get(subject['subject_id'])
                a_series_marks.append(obtained if obtained is not None else 0)
            
            while len(a_series_marks) < 3:
                a_series_marks.append(0)
            
            a_series_avg = sum(a_series_marks) / len(a_series_marks) if a_series_marks else 0
            
            if a_series and a_series[0]['total_marks'] > 0:
                twenty_percent = (a_series_avg / a_series[0]['total_marks']) * 20
            else:
                twenty_percent = 0
            
            if eot_exam and eot_exam.get('total_marks', 0) > 0:
                eot_mark = all_marks.get(eot_exam['id'], {}).get(subject['subject_id'], 0)
                eighty_percent = (eot_mark / eot_exam['total_marks']) * 80
            else:
                eighty_percent = 0
            
            total_percentage = twenty_percent + eighty_percent
            
            identifier = '1' if total_percentage < 40 else ('2' if total_percentage < 70 else '3')
            
            if total_percentage >= 80:
                grade = 'A'
                remark = 'Exceptional'
            elif total_percentage >= 70:
                grade = 'B'
                remark = 'Outstanding'
            elif total_percentage >= 60:
                grade = 'C'
                remark = 'Satisfactory'
            elif total_percentage >= 40:
                grade = 'D'
                remark = 'Basic'
            else:
                grade = 'E'
                remark = 'Insufficient'
            
            subject_results.append({
                'name': subject_name,
                'a1': f"{a_series_marks[0]:.1f}",
                'a2': f"{a_series_marks[1]:.1f}",
                'a3': f"{a_series_marks[2]:.1f}",
                'avg': f"{a_series_avg:.1f}",
                'twenty': f"{twenty_percent:.1f}",
                'eighty': f"{eighty_percent:.1f}",
                'total': f"{total_percentage:.1f}",
                'identifier': identifier,
                'grade': grade,
                'remark': remark
            })
        
        if subject_results:
            overall_percentage = sum([float(s['total']) for s in subject_results]) / len(subject_results)
        else:
            overall_percentage = 0
        
        if overall_percentage >= 80:
            overall_grade = 'A'
            overall_achievement = 'EXCELLENT'
        elif overall_percentage >= 70:
            overall_grade = 'B'
            overall_achievement = 'VERY GOOD'
        elif overall_percentage >= 60:
            overall_grade = 'C'
            overall_achievement = 'GOOD'
        elif overall_percentage >= 40:
            overall_grade = 'D'
            overall_achievement = 'AVERAGE'
        else:
            overall_grade = 'E'
            overall_achievement = 'NEEDS IMPROVEMENT'
        
        overall_ident = '3' if overall_percentage >= 70 else ('2' if overall_percentage >= 40 else '1')
        
        def safe_upper(value, default='N/A'):
            if value is None:
                return default
            return str(value).upper()
        
        def safe_get(dict_obj, key, default=''):
            value = dict_obj.get(key, default)
            if value is None:
                return default
            return value
        
        # Load template
        template_dir = os.path.join(os.path.dirname(__file__), 'templates')
        template_path = os.path.join(template_dir, 'competence_report_template.docx')
        
        if not os.path.exists(template_path):
            os.makedirs(template_dir, exist_ok=True)
            create_default_template(template_path)
        
        doc = DocxTemplate(template_path)
        
        # Prepare context (simplified without images for class reports to improve performance)
        context = {
            'logo': '',
            'student_photo': '',
            'SCHOOL_NAME': safe_upper(institute.get('institute_name'), 'ST CHARLES LWANGA SS-AKASHANDA'),
            'ADDRESS': safe_get(institute, 'address', 'Kampala, Uganda'),
            'PHONE_NUMBER': safe_get(institute, 'phone_number', '+256757632756'),
            'EMAIL': safe_get(institute, 'email', 'info@school.ac.ug'),
            'TERM': safe_upper(term, term),
            'YEAR': str(year),
            'STUDENT_NAME': safe_upper(student.get('name'), 'STUDENT'),
            'GENDER': safe_upper(student.get('gender'), 'NOT SPECIFIED'),
            'CLASS_NAME': safe_upper(student.get('class_name', 'N/A')),
            'SECTION': safe_upper(student.get('category'), 'DAY'),
            'subjects': subject_results,
            'overall_avg': f"{overall_percentage:.1f}",
            'overall_ident': overall_ident,
            'overall_achievement': overall_achievement,
            'overall_grade': overall_grade,
            'teacher_comment': f"{safe_get(student, 'name', 'The student')} is making good progress.",
            'next_term_date': f"05/{datetime.now().year + 1 if datetime.now().month > 8 else datetime.now().month}/2026",
            'printed_date': datetime.now().strftime('%d/%m/%Y'),
        }
        
        # Render template
        doc.render(context)
        
        # Save DOCX to temp file
        docx_path = os.path.join(temp_dir, f"report_{student['id']}_{student['name'].replace(' ', '_')}.docx")
        doc.save(docx_path)
        
        # Convert to PDF
        pdf_path = convert_docx_to_pdf_with_fallback(docx_path, temp_dir, f"report_{student['id']}_{student['name'].replace(' ', '_')}.pdf")
        
        return pdf_path if os.path.exists(pdf_path) else None
        
    except Exception as e:
        print(f"Error generating report for student {student.get('name', 'Unknown')}: {e}")
        return None

def create_zip_archive(pdf_files, class_id, term, year):
    """
    Create a ZIP archive of PDF files as fallback when merging fails
    """
    try:
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for pdf_file in pdf_files:
                arcname = os.path.basename(pdf_file)
                zip_file.write(pdf_file, arcname)
        
        zip_buffer.seek(0)
        
        # Get class name
        class_response = supabase.table('classes')\
            .select('name')\
            .eq('id', class_id)\
            .execute()
        class_name = class_response.data[0]['name'] if class_response.data else f"class_{class_id}"
        
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=f"competence_reports_{class_name}_term{term}_{year}.zip",
            mimetype='application/zip'
        )
    except Exception as e:
        print(f"Error creating ZIP archive: {e}")
        return jsonify({'success': False, 'message': 'Failed to create reports'}), 500

def create_default_template(template_path):
    """Create a default Word template that works properly with docxtpl"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document()
    
    # Set page margins for A4
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # Header table with 3 columns: Logo, School Info, Student Photo
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    
    # Set column widths
    header_table.columns[0].width = Cm(2.5)  # Logo
    header_table.columns[1].width = Cm(10)   # School info
    header_table.columns[2].width = Cm(2.5)  # Student photo
    
    # Remove borders from header table
    for row in header_table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
    
    # Logo cell (left)
    logo_cell = header_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_para.add_run('{{ logo }}')
    run.font.size = Pt(8)
    
    # School info cell (center)
    school_cell = header_table.cell(0, 1)
    school_cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    
    school_name = school_cell.paragraphs[0]
    school_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school_name.add_run('{{ SCHOOL_NAME }}')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 128)
    
    school_address = school_cell.add_paragraph()
    school_address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school_address.add_run('{{ ADDRESS }}').font.size = Pt(10)
    
    school_contact = school_cell.add_paragraph()
    school_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school_contact.add_run('TEL: {{ PHONE_NUMBER }}    EMAIL: {{ EMAIL }}').font.size = Pt(9)
    
    # Student photo cell (right)
    photo_cell = header_table.cell(0, 2)
    photo_cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    photo_para = photo_cell.paragraphs[0]
    photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = photo_para.add_run('{{ student_photo }}')
    run.font.size = Pt(8)
    
    # Add spacing
    doc.add_paragraph()
    
    # Report Title
    report_title = doc.add_paragraph()
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = report_title.add_run('TERM {{ TERM }} REPORT CARD {{ YEAR }}')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(255, 0, 0)
    
    # Student Information Table
    info_table = doc.add_table(rows=2, cols=6)
    info_table.style = 'Table Grid'
    
    # Row 1
    info_table.cell(0, 0).text = 'NAME:'
    info_table.cell(0, 1).text = '{{ STUDENT_NAME }}'
    info_table.cell(0, 2).text = 'GENDER:'
    info_table.cell(0, 3).text = '{{ GENDER }}'
    info_table.cell(0, 4).text = 'TERM:'
    info_table.cell(0, 5).text = '{{ TERM }}'
    
    # Row 2
    info_table.cell(1, 0).text = 'CLASS:'
    info_table.cell(1, 1).text = '{{ CLASS_NAME }}'
    info_table.cell(1, 2).text = 'SECTION:'
    info_table.cell(1, 3).text = '{{ SECTION }}'
    info_table.cell(1, 4).text = 'YEAR:'
    info_table.cell(1, 5).text = '{{ YEAR }}'
    
    # Make header cells bold
    for row in info_table.rows:
        for cell in row.cells:
            if cell.text.strip().endswith(':'):
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        paragraph.runs[0].bold = True
    
    doc.add_paragraph()
    
    # Subjects Table
    subjects_table = doc.add_table(rows=1, cols=11)
    subjects_table.style = 'Table Grid'
    
    # Set column widths
    subjects_table.columns[0].width = Cm(3.5)  # SUBJECT
    for i in range(1, 10):
        subjects_table.columns[i].width = Cm(1.2)  # Number columns
    subjects_table.columns[10].width = Cm(3.5)  # REMARKS
    
    # Table headers
    headers = ['SUBJECT', 'A1', 'A2', 'A3', 'AVG', '20%', '80%', '100%', 'IDENT', 'GRADE', 'REMARKS']
    header_cells = subjects_table.rows[0].cells
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add the Jinja2 for loop start tag
    loop_start_row = subjects_table.add_row()
    loop_start_cell = loop_start_row.cells[0]
    loop_start_cell.text = '{% for subject in subjects %}'
    loop_start_cell.merge(loop_start_row.cells[10])
    loop_start_cell.paragraphs[0].runs[0].font.size = Pt(1)
    loop_start_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Add the actual subject data row
    data_row = subjects_table.add_row()
    data_row.cells[0].text = '{{ subject.name }}'
    data_row.cells[1].text = '{{ subject.a1 }}'
    data_row.cells[2].text = '{{ subject.a2 }}'
    data_row.cells[3].text = '{{ subject.a3 }}'
    data_row.cells[4].text = '{{ subject.avg }}'
    data_row.cells[5].text = '{{ subject.twenty }}'
    data_row.cells[6].text = '{{ subject.eighty }}'
    data_row.cells[7].text = '{{ subject.total }}'
    data_row.cells[8].text = '{{ subject.identifier }}'
    data_row.cells[9].text = '{{ subject.grade }}'
    data_row.cells[10].text = '{{ subject.remark }}'
    
    # Center align the numeric columns
    for i in range(1, 10):
        data_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add the Jinja2 for loop end tag
    loop_end_row = subjects_table.add_row()
    loop_end_cell = loop_end_row.cells[0]
    loop_end_cell.text = '{% endfor %}'
    loop_end_cell.merge(loop_end_row.cells[10])
    loop_end_cell.paragraphs[0].runs[0].font.size = Pt(1)
    loop_end_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph()
    
    # Overall Performance Table
    overall_table = doc.add_table(rows=1, cols=6)
    overall_table.style = 'Table Grid'
    overall_cells = overall_table.rows[0].cells
    overall_cells[0].text = 'OVERALL AVERAGE:'
    overall_cells[1].text = '{{ overall_avg }}%'
    overall_cells[2].text = 'IDENTIFIER:'
    overall_cells[3].text = '{{ overall_ident }}'
    overall_cells[4].text = 'GRADE:'
    overall_cells[5].text = '{{ overall_grade }}'
    
    # Make labels bold
    for i in [0, 2, 4]:
        if overall_cells[i].paragraphs[0].runs:
            overall_cells[i].paragraphs[0].runs[0].bold = True
    
    # Grading Scale
    doc.add_paragraph()
    scale_title = doc.add_paragraph()
    scale_title.add_run('GRADING SCALE').bold = True
    
    scale_table = doc.add_table(rows=2, cols=5)
    scale_table.style = 'Table Grid'
    
    grades = ['A', 'B', 'C', 'D', 'E']
    scores = ['80-100', '70-79', '60-69', '40-59', '0-39']
    
    for i, grade in enumerate(grades):
        scale_table.cell(0, i).text = grade
        scale_table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if scale_table.cell(0, i).paragraphs[0].runs:
            scale_table.cell(0, i).paragraphs[0].runs[0].bold = True
        scale_table.cell(1, i).text = scores[i]
        scale_table.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Comments
    doc.add_paragraph()
    comment_para = doc.add_paragraph()
    comment_para.add_run("CLASS TEACHER'S COMMENT: ").bold = True
    comment_para.add_run("{{ teacher_comment }}")
    
    doc.add_paragraph()
    head_comment = doc.add_paragraph()
    head_comment.add_run("HEADTEACHER'S COMMENT: ").bold = True
    head_comment.add_run("_________________________________________")
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("NEXT TERM BEGINS: {{ next_term_date }}").bold = True
    footer.add_run("    |    ")
    footer.add_run("PRINTED ON: {{ printed_date }}").bold = True
    
    # Save the template
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    doc.save(template_path)
    print(f"✓ Default template created at {template_path}")